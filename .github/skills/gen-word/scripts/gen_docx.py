#!/usr/bin/env python3
"""Generate a professional Word (.docx) document from JSON content data.

Usage:
    python3 gen_docx.py --input content.json --output report.docx --style corporate
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Error: python-docx not installed. Run: pip3 install python-docx", file=sys.stderr)
    sys.exit(1)

STYLES = {
    "corporate": {
        "heading_color": RGBColor(0x1A, 0x36, 0x5D),
        "accent_color": RGBColor(0x31, 0x82, 0xCE),
        "text_color": RGBColor(0x2D, 0x37, 0x48),
        "font_heading": "Calibri",
        "font_body": "Calibri",
        "font_size_title": Pt(24),
        "font_size_h1": Pt(18),
        "font_size_h2": Pt(14),
        "font_size_body": Pt(11),
        "table_header_bg": RGBColor(0x1A, 0x36, 0x5D),
    },
    "academic": {
        "heading_color": RGBColor(0x1A, 0x20, 0x2C),
        "accent_color": RGBColor(0x74, 0x42, 0x10),
        "text_color": RGBColor(0x1A, 0x20, 0x2C),
        "font_heading": "Times New Roman",
        "font_body": "Times New Roman",
        "font_size_title": Pt(22),
        "font_size_h1": Pt(16),
        "font_size_h2": Pt(13),
        "font_size_body": Pt(12),
        "table_header_bg": RGBColor(0x74, 0x42, 0x10),
    },
    "minimal": {
        "heading_color": RGBColor(0x11, 0x18, 0x27),
        "accent_color": RGBColor(0x05, 0x96, 0x69),
        "text_color": RGBColor(0x37, 0x41, 0x51),
        "font_heading": "Arial",
        "font_body": "Arial",
        "font_size_title": Pt(22),
        "font_size_h1": Pt(16),
        "font_size_h2": Pt(13),
        "font_size_body": Pt(11),
        "table_header_bg": RGBColor(0x05, 0x96, 0x69),
    },
}


def apply_run_style(run, font_name, font_size, color=None, bold=False):
    """Apply font style to a run."""
    run.font.name = font_name
    run.font.size = font_size
    if color:
        run.font.color.rgb = color
    run.bold = bold


def add_heading_styled(doc, text, level, style):
    """Add a styled heading paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if level == 0:
        apply_run_style(run, style["font_heading"], style["font_size_title"],
                        style["heading_color"], bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        apply_run_style(run, style["font_heading"], style["font_size_h1"],
                        style["heading_color"], bold=True)
    else:
        apply_run_style(run, style["font_heading"], style["font_size_h2"],
                        style["accent_color"], bold=True)
    para.space_after = Pt(6)
    return para


def add_paragraph_styled(doc, text, style):
    """Add a styled body paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    apply_run_style(run, style["font_body"], style["font_size_body"], style["text_color"])
    para.space_after = Pt(4)
    return para


def add_table_styled(doc, headers, rows, style):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = style["font_size_body"]
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        bg = style["table_header_bg"]
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '{:02X}{:02X}{:02X}'.format(bg[0], bg[1], bg[2]),
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = style["font_size_body"]

    doc.add_paragraph()  # spacing after table


def generate_docx(data, style_name):
    """Generate a Document from structured JSON data."""
    style = STYLES.get(style_name, STYLES["corporate"])
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    title = data.get("title", "Document")
    add_heading_styled(doc, title, 0, style)

    # Subtitle / author
    author = data.get("author", "")
    date = data.get("date", "")
    if author or date:
        meta_text = " | ".join(filter(None, [author, date]))
        para = doc.add_paragraph()
        run = para.add_run(meta_text)
        apply_run_style(run, style["font_body"], Pt(10), style["accent_color"])
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_after = Pt(12)

    # Process sections
    for section_data in data.get("sections", []):
        section_type = section_data.get("type", "text")

        if section_type == "heading":
            level = section_data.get("level", 1)
            add_heading_styled(doc, section_data.get("text", ""), level, style)

        elif section_type == "text":
            add_paragraph_styled(doc, section_data.get("text", ""), style)

        elif section_type == "bullets":
            heading = section_data.get("heading", "")
            if heading:
                add_heading_styled(doc, heading, 2, style)
            for bullet in section_data.get("items", []):
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(bullet)
                apply_run_style(run, style["font_body"], style["font_size_body"],
                                style["text_color"])

        elif section_type == "table":
            heading = section_data.get("heading", "")
            if heading:
                add_heading_styled(doc, heading, 2, style)
            add_table_styled(doc, section_data.get("headers", []),
                             section_data.get("rows", []), style)

        elif section_type == "quote":
            text = section_data.get("text", "")
            author_q = section_data.get("author", "")
            para = doc.add_paragraph()
            run = para.add_run(f'"{text}"')
            run.italic = True
            apply_run_style(run, style["font_body"], style["font_size_body"],
                            style["accent_color"])
            if author_q:
                run2 = para.add_run(f"\n— {author_q}")
                apply_run_style(run2, style["font_body"], Pt(10), style["text_color"])

    return doc


def main():
    parser = argparse.ArgumentParser(
        description="Generate professional Word (.docx) document from JSON data")
    parser.add_argument("--input", required=True, help="Path to JSON file with content data")
    parser.add_argument("--output", required=True, help="Output .docx file path")
    parser.add_argument("--style", choices=list(STYLES.keys()), default="corporate",
                        help="Document style (default: corporate)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = generate_docx(data, args.style)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    size_kb = output_path.stat().st_size / 1024
    section_count = len(data.get("sections", []))
    print(f"✅ Saved: {output_path} ({size_kb:.1f} KB, {section_count} sections, style: {args.style})")


if __name__ == "__main__":
    main()
